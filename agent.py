import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentGenerator:
    """
    A document generator that builds .docx files from structured content.
    """
    def __init__(self):
        self.doc = Document()
        
    def add_title(self, text):
        self.doc.add_heading(text, level=0)
        
    def add_section_heading(self, text, level=1):
        self.doc.add_heading(text, level=level)
        
    def add_paragraph(self, text, style=None, bold=False, italic=False, align=None):
        p = self.doc.add_paragraph(style=style)
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p
        
    def add_bullet_list(self, items):
        for item in items:
            self.doc.add_paragraph(item, style="List Bullet")
            
    def add_table(self, rows_data):
        if not rows_data:
            return
        
        rows = len(rows_data)
        cols = len(rows_data[0])
        table = self.doc.add_table(rows=rows, cols=cols, style="Table Grid")
        
        for i, row_data in enumerate(rows_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = str(cell_data)
                
    def add_image(self, image_path, width_inches=None):
        if width_inches:
            self.doc.add_picture(image_path, width=Inches(width_inches))
        else:
            self.doc.add_picture(image_path)
            
    def add_page_break(self):
        self.doc.add_page_break()
        
    def setup_page_margins(self, top_cm=2.5, bottom_cm=2.5):
        section = self.doc.sections[0]
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)
        
    def set_header(self, text):
        section = self.doc.sections[0]
        header = section.header
        if not header.paragraphs:
            header.add_paragraph(text)
        else:
            header.paragraphs[0].text = text
            
    def save(self, path="output.docx"):
        self.doc.save(path)
        print(f"Document saved to {path}")

if __name__ == "__main__":
    # Example usage based on rapista/skills/docx/SKILL.md
    agent = DocumentGenerator()
    agent.setup_page_margins()
    agent.set_header("Confidential Report")
    
    agent.add_title("Document Generator Report")
    agent.add_paragraph("This is a sample document generated automatically by the Document Generator agent.")
    
    agent.add_section_heading("Features Overview")
    agent.add_bullet_list([
        "Automated docx creation",
        "Structured content generation",
        "Headings, tables, and lists formatting"
    ])
    
    agent.add_section_heading("Data Summary")
    agent.add_table([
        ["Item", "Value", "Status"],
        ["Feature A", "100", "Complete"],
        ["Feature B", "200", "Pending"]
    ])
    
    agent.save("sample_report.docx")
