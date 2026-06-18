from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import datetime

def safe_insert_pBdr(pPr, pBdr):
    """Insert w:pBdr before w:spacing, w:ind, w:jc to satisfy Word schema"""
    for tag in ['w:shd', 'w:tabs', 'w:spacing', 'w:ind', 'w:jc', 'w:rPr']:
        el = pPr.find(qn(tag))
        if el is not None:
            el.addprevious(pBdr)
            return
    pPr.append(pBdr)

def safe_insert_shd(tcPr, shd):
    """Insert w:shd before w:noWrap, w:tcMar to satisfy Word schema"""
    for tag in ['w:noWrap', 'w:tcMar']:
        el = tcPr.find(qn(tag))
        if el is not None:
            el.addprevious(shd)
            return
    tcPr.append(shd)

def set_default_font(doc, name='Calibri', size=Pt(11), color=RGBColor(0x33, 0x33, 0x33)):
    style = doc.styles['Normal']
    font = style.font
    font.name = name
    font.size = size
    font.color.rgb = color
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

def configure_heading_styles(doc):
    for level, size, color_hex in [
        (1, Pt(20), '1A3C5E'),
        (2, Pt(15), '2B579A'),
        (3, Pt(13), '3B6CB4'),
    ]:
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Calibri'
        style.font.size = size
        style.font.bold = True
        style.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
        style.font.space_before = Pt(18 if level == 1 else 12)
        style.font.space_after = Pt(8)

def add_section_break(doc, orientation=WD_SECTION_START.NEW_PAGE):
    section = doc.add_section(orientation)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    return section

def add_header_rule(section):
    p = section.header.paragraphs[0]
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    safe_insert_pBdr(pPr, pBdr)

def set_header_content(section, left_text=None, right_text=None):
    add_header_rule(section)
    p = section.header.paragraphs[0]
    p.clear()

    if left_text:
        run = p.add_run(left_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = 'Calibri'

    if right_text:
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
        run2 = p.add_run(f'\t{right_text}')
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run2.font.name = 'Calibri'

def set_footer_content(section, left_text=None, center_text=None):
    p = section.footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if left_text:
        run = p.add_run(left_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = 'Calibri'
        p.add_run('\t')

def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run(' of ')

    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run5._r.append(fldChar3)

    run6 = paragraph.add_run()
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = ' NUMPAGES '
    run6._r.append(instrText2)

    run7 = paragraph.add_run()
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run7._r.append(fldChar4)

def set_page_number(section, style='DECIMAL', start=1):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        # Word Schema requires pgNumType before cols, docGrid
        cols = sectPr.find(qn('w:cols'))
        docGrid = sectPr.find(qn('w:docGrid'))
        if cols is not None:
            cols.addprevious(pgNumType)
        elif docGrid is not None:
            docGrid.addprevious(pgNumType)
        else:
            sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), style)
    pgNumType.set(qn('w:start'), str(start))

def add_cover_page(doc, title, subtitle=None, author=None, date=None, color_hex='1A3C5E'):
    for _ in range(6):
        doc.add_paragraph()

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    safe_insert_pBdr(pPr, pBdr)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    run = p_title.add_run(title)
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    run.font.name = 'Calibri'

    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sub.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = 'Calibri'

    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = p_line2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '18')
    top.set(qn('w:space'), '6')
    top.set(qn('w:color'), color_hex)
    pBdr2.append(top)
    safe_insert_pBdr(pPr2, pBdr2)

    for _ in range(4):
        doc.add_paragraph()

    meta_items = []
    if author:
        meta_items.append(('Prepared by', author))
    if date:
        meta_items.append(('Date', date))
    else:
        meta_items.append(('Date', datetime.date.today().strftime('%B %d, %Y')))

    for label, value in meta_items:
        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p_meta.add_run(f'{label}:  ')
        run_l.font.size = Pt(10)
        run_l.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_l.font.name = 'Calibri'
        run_v = p_meta.add_run(value)
        run_v.font.size = Pt(11)
        run_v.font.bold = True
        run_v.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
        run_v.font.name = 'Calibri'

def add_toc(doc, title='Table of Contents'):
    doc.add_paragraph()
    h = doc.add_heading(title, level=1)

    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = p.add_run('[Right-click → Update Field to generate Table of Contents]')
    run4.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run4.font.size = Pt(10)
    run4.font.italic = True

    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)

    doc.add_paragraph()

# ── Build document ──
doc = Document()

# ── Configure defaults ──
set_default_font(doc)
configure_heading_styles(doc)

# ── Section 0: Cover page (no header/footer) ──
section0 = doc.sections[0]
section0.header.is_linked_to_previous = False
section0.footer.is_linked_to_previous = False
section0.header.paragraphs[0].clear()
section0.footer.paragraphs[0].clear()
set_page_number(section0, 'upperRoman', 1)

add_cover_page(
    doc,
    title='Steins;Gate',
    subtitle='A Comprehensive Research Report on Time Travel Mechanics and Narrative Impact',
    author='AI Research Analyst',
    date=datetime.date.today().strftime('%B %d, %Y'),
)

# ── Section 1: TOC + Body ──
section1 = add_section_break(doc)
set_page_number(section1, 'decimal', 1)
set_header_content(section1, left_text='Steins;Gate Research Report', right_text='Page ')
add_page_number_field(section1.header.paragraphs[0])
set_footer_content(section1, left_text='Confidential Research')

# TOC
add_toc(doc)
doc.add_page_break()

# ── Content sections ──
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'Steins;Gate is a critically acclaimed science fiction anime that explores the profound consequences '
    'of time travel. Adapted from the visual novel by 5pb. and Nitroplus, it masterfully blends complex '
    'theoretical physics concepts with emotional character development. This report provides an in-depth '
    'analysis of its time travel mechanics, central themes, and cultural impact within the anime medium.'
)

doc.add_heading('Core Time Travel Mechanics', level=1)
doc.add_paragraph(
    'The temporal mechanics of Steins;Gate are grounded in real-world hypotheses, specifically branching '
    'universe models and concepts of causality. The series uses these mechanics consistently, resulting in '
    'one of the most logically sound depictions of time travel in fiction.'
)

doc.add_heading('1. The Phone Microwave (Name Subject to Change)', level=2)
doc.add_paragraph(
    'The initial method of temporal manipulation discovered by the Future Gadget Laboratory. By connecting '
    'a microwave to a mobile phone, the characters are able to send text messages (D-Mails) into the past. '
    'This method relies on manipulating miniature black holes to compress data.'
)

doc.add_heading('2. D-Mails and Butterfly Effects', level=2)
doc.add_paragraph(
    'Sending a D-Mail alters the past, causing the timeline to shift into a new "Attractor Field." The extent '
    'of the shift depends on how the recipient acts upon the received message. This perfectly illustrates the '
    'Butterfly Effect, where seemingly insignificant changes cause massive ripple effects on the world line.'
)

doc.add_heading('3. The Time Leap Machine', level=2)
doc.add_paragraph(
    "An advancement of the Phone Microwave, the Time Leap Machine digitizes the user's memories and sends "
    "them back to their past self. Unlike D-Mails, this allows the user to directly experience and alter "
    "events while retaining knowledge of future outcomes."
)

doc.add_heading('Attractor Fields and World Lines', level=1)
doc.add_paragraph(
    'The universe of Steins;Gate is organized into "Attractor Fields," which are bundles of similar world lines '
    'that converge on a predetermined event (e.g., the death of a specific character). Altering a world line '
    'within an Attractor Field will not change the convergent event. To escape the convergent event, a shift '
    'significant enough to cross the 1% divergence threshold into a different Attractor Field is required.'
)

# Professional table
doc.add_heading('Key Attractor Fields', level=2)
table = doc.add_table(rows=4, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Attractor Field', 'Divergence Range', 'Defining Characteristics']
for i, hdr in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = hdr
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A3C5E" w:val="clear"/>')
    safe_insert_shd(cell._tc.get_or_add_tcPr(), shading)

attractor_fields = [
    ('Alpha', '0.000000% - 0.999999%', 'SERN establishes a dystopia. Mayuri Shiina dies.'),
    ('Beta', '1.000000% - 1.999999%', 'World War III occurs. Makise Kurisu dies.'),
    ('Steins Gate', '1.048596%', 'The ideal timeline where neither tragedy occurs.'),
]
for r, (af, dr, dc) in enumerate(attractor_fields, 1):
    table.rows[r].cells[0].text = af
    table.rows[r].cells[1].text = dr
    table.rows[r].cells[2].text = dc
    if r % 2 == 0:
        for cell in table.rows[r].cells:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F5FA" w:val="clear"/>')
            safe_insert_shd(cell._tc.get_or_add_tcPr(), shading)

doc.add_heading('Reading Steiner', level=1)
doc.add_paragraph(
    'The protagonist, Rintarou Okabe, possesses an ability known as "Reading Steiner." When a world line '
    'shifts, Okabe retains his memories from the previous timeline, unlike the rest of humanity, whose memories '
    'are overwritten to match the new reality. This ability is essential for observing and correcting temporal '
    'changes but places an immense psychological burden on Okabe.'
)

doc.add_heading('Thematic Significance', level=1)
doc.add_paragraph('The series explores several profound themes through its sci-fi premise:')

themes = [
    'The burden of knowledge and responsibility.',
    'The conflict between determinism and free will.',
    'The psychological trauma associated with repeated failure.',
    'The strength of human connection and sacrifice across timelines.',
]
for theme in themes:
    doc.add_paragraph(theme, style='List Bullet')

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'Steins;Gate stands as a pinnacle of time travel storytelling. By meticulously establishing '
    'its rules and consistently adhering to them, it avoids the common plot holes associated with the '
    'genre. Its character-driven narrative ensures that the high-stakes science fiction elements are grounded '
    'in profound emotional stakes, making it a masterclass in both anime and speculative fiction.'
)

# ── Save ──
doc.save('steins_gate_research_report.docx')
